import pandas as pd
import logging
from pathlib import Path
from docx import Document


def get_docx_report_for_all_UPD(df: pd.DataFrame, target_dir: Path):

    # 1. Список уникальных номеров УПД
    numbers = df["Счет-фактура №"].dropna().unique().tolist()
    # 2. Список уникальных адресов
    consignees = df["Грузополучатель и его адрес:"].dropna().unique().tolist()
    # 3. Общая сумма по всем документам
    # errors='coerce' превратит мусор в NaN, чтобы sum() не сломался
    payment_sum = pd.to_numeric(df["Всего к оплате (9)"], errors='coerce').sum()


    payment_sum = f"{payment_sum:,.2f}"
    logging.info(f"Количество номеров: {len(numbers)}")
    logging.info(f"Количество грузополучателей: {len(consignees)}")

    numbers_formatted = ", ".join(sorted(numbers))
    
    logging.info(f"Номера: {numbers_formatted}")
    logging.info(f"Грузополучатели: {consignees}")
    logging.info(f"Общая сумма: {payment_sum}")

    # write to docx file
    logging.info("Создаю документ docx...")
    document = Document()

    logging.info("Пишу параграф с общей суммой...")
    document.add_paragraph(f"Общая сумма: {payment_sum}")

    logging.info("Пишу параграф с номерами...")
    document.add_paragraph(f"Номера: {numbers_formatted}")

    logging.info("Пишу параграф с грузополучателями...")
    document.add_paragraph("Грузополучатели:")
    for c in consignees:
        document.add_paragraph(f"– {c}")
    
    logging.info("Сохраняю документ...")
    
    output_dir = target_dir / "output"
    output_dir.mkdir(exist_ok=True)
    document.save(output_dir / "report.docx")