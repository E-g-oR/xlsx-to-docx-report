import logging
from pathlib import Path
from setup_logs import setup_logs
from docx import Document
import tkinter as tk
from tkinter import filedialog
import sys
import cProfile
import pstats
from index_data import index_data
import pandas as pd

def get_working_directory():
    """Открывает окно выбора папки."""
    root = tk.Tk()
    root.withdraw()  # Скрываем основное маленькое окно tkinter
    root.attributes('-topmost', True)  # Поверх всех окон
    
    selected_dir = filedialog.askdirectory(title="Выберите папку с файлами Excel")
    
    if not selected_dir:
        print("Папка не выбрана. Выход...")
        sys.exit(0)
        
    return Path(selected_dir)


def main():
    target_dir = get_working_directory()
    # target_dir = Path("./assets")

    setup_logs(target_dir)
    logging.info("Start")
    try:
        logging.info("Processing started")
        logging.info("Выбираю файлы .xlsx, начинающиеся с \"УПД\":")

        df = index_data(target_dir)

        # 1. Список уникальных номеров УПД
        numbers = df["Счет-фактура №"].dropna().unique().tolist()
        # 2. Список уникальных адресов
        consignees = df["Грузополучатель и его адрес:"].dropna().unique().tolist()
        # 3. Общая сумма по всем документам
        # errors='coerce' превратит мусор в NaN, чтобы sum() не сломался
        payment_sum = pd.to_numeric(df["Всего к оплате (9)"], errors='coerce').sum()


        payment_sum = f"{payment_sum:,.2f}"
        logging.info(f"Количество номеров: {len(numbers)}")
        logging.info(f"Количество грузополучателей: {len(consignees)}")

        numbers_formatted = ", ".join(sorted(numbers))
        
        logging.info(f"Номера: {numbers_formatted}")
        logging.info(f"Грузополучатели: {consignees}")
        logging.info(f"Общая сумма: {payment_sum}")

        # write to docx file
        logging.info("Создаю документ docx...")
        document = Document()

        logging.info("Пишу параграф с общей суммой...")
        document.add_paragraph(f"Общая сумма: {payment_sum}")

        logging.info("Пишу параграф с номерами...")
        document.add_paragraph(f"Номера: {numbers_formatted}")

        logging.info("Пишу параграф с грузополучателями...")
        document.add_paragraph("Грузополучатели:")
        for c in consignees:
            document.add_paragraph(f"– {c}")
        
        logging.info("Сохраняю документ...")
        
        output_dir = target_dir / "output"
        output_dir.mkdir(exist_ok=True)
        document.save(output_dir / "report.docx")

        logging.info("Готово.")
    
    except Exception as e:
            logging.exception("Fatal error: %s", e)  # автоматически пишет traceback
            raise

if __name__ == "__main__":
    with cProfile.Profile() as pr:
        main()
    
    # Обрабатываем результаты
stats = pstats.Stats(pr)
stats.sort_stats('cumtime') # Сортируем по кумулятивному времени

# ОГРАНИЧИВАЕМ ВЫВОД:
stats.print_stats(10) # Выведет только ТОП-10 самых "тяжелых" функций



